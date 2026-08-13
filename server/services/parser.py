"""JD 文本解析与简历文件文本提取。

全部本地规则实现，不联网、不调用外部服务。
解析结果作为初稿，用户可在界面上手动修正；也可用 AI 通道覆盖。
"""

import re
from pathlib import Path

CITIES = [
    "北京", "上海", "广州", "深圳", "杭州", "南京", "成都", "武汉", "西安", "苏州",
    "厦门", "长沙", "重庆", "天津", "青岛", "郑州", "合肥", "无锡", "宁波", "东莞",
    "佛山", "珠海", "福州", "济南", "大连", "沈阳", "昆明", "贵阳", "南昌", "远程",
]

# 三个目标方向的技能关键词库，用于匹配打分
KEYWORD_LIB = {
    "销售": [
        "大客户", "KA", "渠道", "商务拓展", "BD", "线索", "leads", "转化率", "客单价",
        "回款", "签单", "续约", "复购", "客户成功", "SaaS", "To B", "ToB", "招投标",
        "合同", "报价", "谈判", "拜访", "CRM", "销售漏斗", "配额", "quota", "业绩",
        "代理商", "经销商", "客户关系", "解决方案", "售前", "行业客户",
    ],
    "运营": [
        "用户运营", "内容运营", "活动运营", "社区运营", "私域", "社群", "增长",
        "拉新", "留存", "促活", "转化", "复购", "GMV", "DAU", "MAU", "留存率",
        "AB测试", "A/B", "数据分析", "SQL", "看板", "策略运营", "商家运营",
        "品类运营", "供给", "履约", "投放", "ROI", "选品", "达人", "KOL", "MCN",
        "直播", "电商运营", "内容生态", "标签体系", "用户分层", "RFM", "SOP",
    ],
    "市场": [
        "品牌", "整合营销", "IMC", "投放", "买量", "信息流", "SEM", "SEO", "ASO",
        "公关", "PR", "媒介", "KOL", "达人", "种草", "小红书", "抖音", "微信生态",
        "内容营销", "案例", "白皮书", "线索", "MQL", "SQL线索", "获客成本", "CAC",
        "ROI", "转化漏斗", "活动策划", "发布会", "campaign", "预算", "创意",
        "私域", "会员", "用户洞察", "市场调研", "竞品分析", "GTM",
    ],
    "通用": [
        "跨部门", "项目管理", "OKR", "KPI", "复盘", "汇报", "英语", "团队管理",
        "带团队", "从0到1", "0-1", "数据驱动", "商业化", "预算管理", "流程优化",
        "PPT", "Excel", "Python", "SQL", "BI", "Tableau", "飞书", "钉钉",
    ],
}

DUTY_HEADS = [
    "岗位职责", "工作职责", "职位描述", "工作内容", "职责描述", "你将负责",
    "主要职责", "岗位描述", "job description", "responsibilities",
]
REQ_HEADS = [
    "任职要求", "岗位要求", "任职资格", "职位要求", "我们希望你", "要求",
    "任职条件", "加分项", "requirements", "qualifications",
]


def _split_lines(text: str) -> list[str]:
    text = text.replace("\r", "\n").replace("•", "\n").replace("·", "\n")
    lines = [re.sub(r"^[\s\-\*\d\.、）\)]+", "", ln).strip() for ln in text.split("\n")]
    return [ln for ln in lines if ln]


def _hit_head(line: str, heads: list[str]) -> bool:
    low = line.lower()
    return any(h.lower() in low for h in heads) and len(line) < 30


def parse_jd(text: str) -> dict:
    """把 JD 原文拆成结构化字段。返回值可直接存 jd.parsed_json。"""
    result = {
        "responsibilities": [],
        "requirements": [],
        "keywords": [],
        "years_required": None,
        "education_required": None,
        "highlights": [],
    }
    if not text or not text.strip():
        return result

    lines = _split_lines(text)
    section = None
    for line in lines:
        if _hit_head(line, DUTY_HEADS):
            section = "responsibilities"
            continue
        if _hit_head(line, REQ_HEADS):
            section = "requirements"
            continue
        if section and 4 <= len(line) <= 200:
            result[section].append(line)

    # 没有明显分段时，整体当作要求处理，保证关键词能抽出来
    if not result["responsibilities"] and not result["requirements"]:
        result["requirements"] = [ln for ln in lines if 4 <= len(ln) <= 200]

    years = re.search(r"(\d+)\s*年以上|(\d+)\s*年\+|(\d+)\s*-\s*\d+\s*年", text)
    if years:
        result["years_required"] = int(next(g for g in years.groups() if g))
    for edu in ["博士", "硕士", "研究生", "本科", "大专", "统招本科"]:
        if edu in text:
            result["education_required"] = edu
            break

    result["keywords"] = extract_keywords(text)
    return result


def extract_keywords(text: str) -> list[str]:
    """按关键词库抽取命中的技能词，保持原库大小写。"""
    low = text.lower()
    hits: list[str] = []
    for words in KEYWORD_LIB.values():
        for word in words:
            if word.lower() in low and word not in hits:
                hits.append(word)
    return hits


def guess_meta(text: str) -> dict:
    """尝试猜出公司、岗位、薪资、地点，作为表单默认值。"""
    lines = _split_lines(text)
    meta = {"company": "", "title": "", "salary_range": "", "location": ""}

    salary = re.search(r"(\d+\s*[-~]\s*\d+\s*[kK])|(\d+\s*[kK]\s*[-~]\s*\d+\s*[kK])"
                       r"|(\d+\s*[-~]\s*\d+\s*万)|面议", text)
    if salary:
        meta["salary_range"] = salary.group(0).replace(" ", "")

    for city in CITIES:
        if city in text:
            meta["location"] = city
            break

    # ---- 公司名识别（多层策略）----
    # 1) 显式标签：公司名/公司：/公司名称/【公司】
    for line in lines[:20]:
        m = re.search(r"(?:公司名[称]?|公司)\s*[:：]\s*(.{2,30})", line)
        if m:
            meta["company"] = m.group(1).strip()
            break
    # 2) 常见公司后缀（科技/集团/有限/网络/信息/股份/文化/传媒/咨询/教育/数据/智能/软件/互联网/电子/贸易/实业）
    if not meta["company"]:
        for line in lines[:20]:
            if len(line) <= 30 and re.search(
                r"(公司|集团|科技|网络|信息|股份|文化|传媒|咨询|教育|数据|智能|软件|互联网|电子|贸易|实业|电商)",
                line
            ) and not re.search(r"(岗位|职责|要求|任职|学历|经验|薪资|福利|周末|五险|负责|本科|大专|硕士)", line):
                meta["company"] = line
                break
    # 3) 首行短句启发：第一行短、无岗位动词、无标点，可能是公司名
    if not meta["company"] and lines:
        first = lines[0]
        if len(first) <= 12 and not re.search(r"(运营|销售|市场|品牌|经理|总监|主管|工程师|设计师|顾问|专员|助理)", first):
            # 如果第二行是明显岗位名，则首行大概率是公司
            if len(lines) > 1 and re.search(r"(运营|销售|市场|品牌|经理|总监|主管|工程师|设计师|顾问|专员|助理|管培)", lines[1]):
                meta["company"] = first

    # ---- 岗位名识别 ----
    if not meta["title"]:
        # 显式标签
        for line in lines[:20]:
            m = re.search(r"(?:岗位|职位|岗位名[称]?|职位名[称]?)\s*[:：]\s*(.{2,30})", line)
            if m:
                meta["title"] = m.group(1).strip()
                break
    if not meta["title"]:
        for line in lines[:8]:
            if len(line) <= 30 and re.search(
                r"(运营|销售|市场|品牌|增长|营销|商务|BD|渠道|经理|专家|总监|主管|工程师|设计师|顾问|专员|助理|管培|策划|编辑|主播|达人|讲师)",
                line
            ) and not re.search(r"(岗位职责|工作职责|任职要求|我们希望你|负责|职责描述)", line):
                meta["title"] = line
                break

    # 清理公司名里可能夹带的杂质
    if meta["company"]:
        meta["company"] = meta["company"].strip("【】[]（）() \t")
    if meta["title"]:
        meta["title"] = meta["title"].strip("【】[]（）() \t")
    return meta


def guess_direction(text: str) -> str:
    """按关键词命中数判断这条 JD 更偏哪个方向。"""
    low = text.lower()
    scores = {}
    for name in ["销售", "运营", "市场"]:
        scores[name] = sum(1 for w in KEYWORD_LIB[name] if w.lower() in low)
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "其他"


# ---------------- 简历文件文本提取 ----------------
def extract_resume_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            import docx

            doc = docx.Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(p for p in parts if p and p.strip())
        if suffix == ".pdf":
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # 解析失败不阻断上传
        return f"[自动解析失败，可手动粘贴简历要点] {exc}"
    return ""
